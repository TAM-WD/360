import os
import json
import time
import logging
import requests
import tempfile
import subprocess
import struct
import re
import imageio_ffmpeg
from datetime import datetime
from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ============================================================
# Конфигурация
# ============================================================

MESSENGER_BOT_TOKEN = os.environ["MESSENGER_BOT_TOKEN"]
MESSENGER_BASE_URL  = "https://botapi.messenger.yandex.net/bot/v1"

YC_API_KEY   = os.environ["YC_API_KEY"]
YC_FOLDER_ID = os.environ["YC_FOLDER_ID"]

SPEECHKIT_V1_URL = "https://stt.api.cloud.yandex.net/speech/v1/stt:recognize"

YANDEXGPT_COMPLETION_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
YANDEXGPT_MODEL_URI      = f"gpt://{YC_FOLDER_ID}/yandexgpt/latest"

FFMPEG_PATH = imageio_ffmpeg.get_ffmpeg_exe()

MAX_V1_SIZE = 1 * 1024 * 1024

AUDIO_EXTENSIONS = {
    ".ogg", ".opus", ".mp3", ".wav", ".m4a", ".flac",
    ".amr", ".aac", ".wma", ".webm", ".spx",
}

# ============================================================
# СТРУКТУРА ЧЕК-ЛИСТА (21 блок) — данные для шаблона
# ============================================================

CHECKLIST_BLOCKS = [
    {
        "number": 1,
        "title": "Обучение сотрудников",
        "items": [
            "1.1 Проводится ли обучение для каждого сотрудника, назначенного на новую должность?",
            "1.2 Включает ли первоначальное обучение подробный разбор возможных производственных рисков и опасностей?",
            "1.3 Проводится ли достаточно наглядный инструктаж по использованию СИЗ?",
            "1.4 Осуществляется ли обучение по использованию аварийно-спасательного оборудования?",
        ],
    },
    {
        "number": 2,
        "title": "Окружающая среда",
        "items": [
            "2.1 Имеются ли средства для работы в очень жарких/холодных условиях?",
            "2.2 Водоотталкивающие/зимние СИЗ удобны, не создают рисков при ношении?",
            "2.3 Безопасны ли рабочие и опорные поверхности при намокании?",
            "2.4 Знают ли работники симптомы теплового удара, обморожения, гипотермии?",
        ],
    },
    {
        "number": 3,
        "title": "Освещённость",
        "items": [
            "3.1 Достаточен ли уровень освещения для безопасной и комфортной работы?",
            "3.2 Вызывает ли освещение блики на рабочих поверхностях, экранах?",
            "3.3 Достаточен ли аварийный свет и проходит ли он регулярную проверку?",
        ],
    },
    {
        "number": 4,
        "title": "Проемы/отверстия в полу и стенах",
        "items": [
            "4.1 Оснащены ли проёмы и двери ограждениями?",
            "4.2 Ограждены ли временные проёмы или находятся ли поблизости сотрудники для контроля?",
        ],
    },
    {
        "number": 5,
        "title": "Лестницы, стремянки и платформы",
        "items": [
            "5.1 Лестницы и перила находятся в хорошем состоянии?",
            "5.2 Лестницы не имеют визуальных дефектов и повреждений?",
            "5.3 Лестницы/стремянки правильно установлены до начала работы?",
            "5.4 Надлежащим ли образом закреплены приподнятые платформы, есть ли поручни?",
        ],
    },
    {
        "number": 6,
        "title": "Подъемные устройства",
        "items": [
            "6.1 Подъемные устройства используются только в рамках своей грузоподъемности?",
            "6.2 Указаны ли лимиты по грузоподъемности на оборудовании?",
            "6.3 Проходит ли оборудование регулярный технический осмотр и обслуживание?",
            "6.4 Операторы обучены работе с подъемным оборудованием?",
        ],
    },
    {
        "number": 7,
        "title": "Ограниченные пространства",
        "items": [
            "7.1 Доступны и соблюдаются ли процедуры по работе в замкнутых пространствах?",
            "7.2 Достаточны ли процедуры входа и выхода из ограниченных пространств?",
            "7.3 Внедрены ли аварийные и спасательные процедуры при ЧП в ограниченных пространствах?",
        ],
    },
    {
        "number": 8,
        "title": "Хозяйственная часть (АБП)",
        "items": [
            "8.1 В помещениях поддерживается чистота и порядок?",
            "8.2 Нет ли на полах торчащих гвоздей, заноз, дыр, незакрепленных досок?",
            "8.3 Проходы и коридоры не загромождены?",
            "8.4 Чётко ли обозначены постоянные проходы и коридоры?",
            "8.5 Оборудованы ли открытые ямы/проемы, резервуары крышками/перилами?",
        ],
    },
    {
        "number": 9,
        "title": "Электробезопасность",
        "items": [
            "9.1 Соблюдаются ли стандарты при эксплуатации и обслуживании электрооборудования?",
            "9.2 Всё оборудование заземлено правильно?",
            "9.3 Ручной электроинструмент заземлен или с двойной изоляцией?",
            "9.4 Распределительные коробки закрыты?",
            "9.5 Удлинители не находятся в проходах, где возможны повреждения?",
            "9.6 Используется ли стационарная проводка вместо удлинителей?",
        ],
    },
    {
        "number": 10,
        "title": "Инструменты и оборудование",
        "items": [
            "10.1 Инструкции производителя хранятся и используются при работе?",
            "10.2 Электроинструмент соответствует стандартам предприятия?",
            "10.3 Бракованный инструмент маркируется и выводится из эксплуатации?",
            "10.4 Организовано ли обучение по безопасному использованию инструментов?",
        ],
    },
    {
        "number": 11,
        "title": "Ограждения оборудования",
        "items": [
            "11.1 Всё движущееся оборудование защищено ограждениями?",
            "11.2 Ограждения соответствуют стандартам безопасности?",
            "11.3 Все ограждения установлены и выполняют своё назначение?",
            "11.4 Соблюдаются ли процедуры блокировки при обслуживании без ограждений?",
        ],
    },
    {
        "number": 12,
        "title": "Уровень шума",
        "items": [
            "12.1 Проводятся ли регулярные замеры уровня шума?",
            "12.2 СИЗ для защиты слуха предоставлены и используются правильно?",
        ],
    },
    {
        "number": 13,
        "title": "Временные рабочие конструкции",
        "items": [
            "13.1 Временные конструкции используются только при невозможности постоянных?",
            "13.2 Надлежащим ли образом укреплены выемки/ямы/котлованы?",
        ],
    },
    {
        "number": 14,
        "title": "Средства индивидуальной защиты (СИЗ)",
        "items": [
            "14.1 Все необходимые СИЗ предоставлены, хранение правильное, использование по назначению?",
            "14.2 Выбор СИЗ выполнен с учетом конкретных опасных факторов?",
            "14.3 Используемые СИЗ надёжны?",
            "14.4 Обозначены ли опасные зоны, требующие ношения СИЗ?",
        ],
    },
    {
        "number": 15,
        "title": "Обращение с материалами и их хранение",
        "items": [
            "15.1 Обеспечено ли безопасное расстояние для проезда техники?",
            "15.2 Материалы складированы устойчивым и надёжным способом?",
            "15.3 В зоне хранения нет рисков опрокидывания?",
            "15.4 К управлению погрузчиком допущены только обученные специалисты?",
            "15.5 Зарядка аккумуляторов — только в отведённых местах?",
            "15.6 На ж/д путях применяются необходимые предупреждающие знаки?",
            "15.7 Указаны ли нормы допустимых нагрузок на стеллажи, полы?",
            "15.8 Стеллажи нагружаются только в пределах максимальной грузоподъемности?",
            "15.9 Тали, канаты и стропы соответствуют требованиям по грузоподъемности?",
            "15.10 Стропы проходят ежедневный осмотр перед применением?",
            "15.11 Используются только проверенные поддоны и платформы?",
            "15.12 Персонал поднимает грузы правильной техникой?",
        ],
    },
    {
        "number": 16,
        "title": "Опасные продукты и химикаты",
        "items": [
            "16.1 Паспорт безопасности изучен до взаимодействия с опасными веществами?",
            "16.2 Для работы с опасными продуктами используются соответствующие СИЗ?",
            "16.3 Хранение не допускает совместимости несочетаемых продуктов?",
            "16.4 Опасные вещества размещены в отдалении от источников тепла?",
            "16.5 Контейнеры проверяются на протечки, повреждения?",
            "16.6 Используются ли поддоны под ёмкостями?",
            "16.7 Вся опасная продукция промаркирована?",
            "16.8 Средства для ликвидации разливов доступны?",
            "16.9 Для горючих веществ предусмотрены устройства заземления?",
            "16.10 Вентиляция зоны хранения соответствует требованиям?",
        ],
    },
    {
        "number": 17,
        "title": "Технологические процессы",
        "items": [
            "17.1 Повторяющиеся действия на рабочем месте оптимизированы?",
            "17.2 Паспорта безопасности доступны для всех сотрудников?",
            "17.3 Риски и опасности отмечены предупреждающими табличками?",
            "17.4 Техника проходит обязательный осмотр перед сменой и периодическое ТО?",
            "17.5 Процедуры LOTO внедрены и строго соблюдаются?",
            "17.6 Вентиляционное оборудование исправно и работает эффективно?",
            "17.7 Система вытяжки/пылеулавливания в исправном состоянии?",
            "17.8 Души/станции самоспасения и фонтанчики для глаз доступны и исправны?",
        ],
    },
    {
        "number": 18,
        "title": "Бытовые помещения для сотрудников",
        "items": [
            "18.1 Помещения содержатся в чистоте и надлежащей санитарии?",
            "18.2 Мебель, шкафчики, сантехника в исправном состоянии?",
            "18.3 Столовая отделена от зоны хранения опасных веществ?",
            "18.4 Помещения оборудованы исправными мойками с гигиеническими средствами?",
        ],
    },
    {
        "number": 19,
        "title": "Пожарная безопасность",
        "items": [
            "19.1 На каждом рабочем месте размещён актуальный план эвакуации?",
            "19.2 Все сотрудники знают этот план?",
            "19.3 Огнетушители подобраны с учетом масштабов возможных пожаров?",
            "19.4 Количество огнетушителей достаточное?",
            "19.5 Места расположения огнетушителей визуально отмечены?",
            "19.6 Огнетушители правильно установлены и легкодоступны?",
            "19.7 Все огнетушители не просроченные, заряжены и готовы?",
            "19.8 Огнетушители специального назначения промаркированы?",
        ],
    },
    {
        "number": 20,
        "title": "Пути эвакуации и эвакуационные выходы",
        "items": [
            "20.1 Достаточно ли выходов для быстрой эвакуации?",
            "20.2 Доступ к выходам свободен, пути не загромождены?",
            "20.3 Двери открываются в экстренной ситуации беспрепятственно?",
            "20.4 Все эвакуационные выходы чётко промаркированы?",
            "20.5 На выходах и путях эвакуации есть аварийное освещение?",
            "20.6 Проходы и выходы эвакуации свободны?",
        ],
    },
    {
        "number": 21,
        "title": "Медицинская помощь и первая помощь",
        "items": [
            "21.1 Все сотрудники знают, как получить первую помощь?",
            "21.2 Лица, оказывающие первую помощь, знают куда доставлять пострадавшего?",
            "21.3 На каждой смене есть обученные оказанию первой помощи сотрудники?",
            "21.4 Аптечки полностью укомплектованы, ЛС не просрочены?",
            "21.5 Медикаменты пополняются своевременно по мере расходования?",
        ],
    },
]

CHECKLIST_STRUCTURE = "\n\n".join(
    f"БЛОК {b['number']}. {b['title']}\n" + "\n".join(b["items"])
    for b in CHECKLIST_BLOCKS
)

# ============================================================
# СИСТЕМНЫЕ ПРОМПТЫ
# ============================================================

SYSTEM_PROMPT_EXTRACT = """Ты — ассистент по охране труда и промышленной безопасности.

Тебе приходит расшифровка голосового сообщения сотрудника, описывающего обход объекта.

Извлеки ВСЮ информацию и верни строго валидный JSON (без markdown):

{
  "object_name": "название объекта или пустая строка",
  "address": "адрес или пустая строка",
  "inspector": "имя проверяющего или пустая строка",
  "director": "имя директора/управляющего или пустая строка",
  "employees_on_shift": "сотрудники на смене или пустая строка",
  "notes": "общие замечания",
  "findings": [
    {
      "item_id": "номер пункта (например 19.3), или пустая строка если не определить точно",
      "block_number": числовой номер блока от 1 до 21,
      "description": "что обнаружено",
      "status": "ok | violation | warning | not_checked",
      "comment": "детали",
      "deadline": "срок устранения или пустая строка"
    }
  ]
}

Правила:
- Извлекай ВСЮ информацию.
- «Всё нормально» = status: "ok".
- Проблема = "violation" или "warning".
- Соотноси с номерами блоков (1-21) максимально точно.
- ТОЛЬКО валидный JSON, без комментариев и markdown."""

SYSTEM_PROMPT_FILL = f"""Ты — ассистент по охране труда.

Тебе приходят данные из голосового отчёта проверяющего (JSON).
Твоя задача — для КАЖДОГО пункта чек-листа определить оценку.

Структура чек-листа:
{CHECKLIST_STRUCTURE}

Ответь строго в формате JSON (без markdown). Массив объектов:

[
  {{
    "item_id": "1.1",
    "assessment": "ДА | НЕТ | Н/П | НЕ ПРОВЕРЕНО",
    "comment": "комментарий или пустая строка",
    "deadline": "срок устранения или пустая строка"
  }},
  ...
]

Правила:
- Для КАЖДОГО пункта от 1.1 до 21.5 должна быть запись.
- Если есть данные — заполни оценку и комментарий.
- Если данных нет — «НЕ ПРОВЕРЕНО», пустой комментарий.
- ТОЛЬКО валидный JSON-массив."""

SYSTEM_PROMPT_SUMMARY = """Ты — ассистент по охране труда.

Тебе приходит заполненный чек-лист. Сформируй КРАТКИЙ ИТОГОВЫЙ АКТ:

1. ШАПКА: объект, дата, проверяющий
2. КРИТИЧЕСКИЕ НАРУШЕНИЯ — список с номерами пунктов
3. ПРЕДУПРЕЖДЕНИЯ — список
4. РЕКОМЕНДАЦИИ ПО УСТРАНЕНИЮ — действия и сроки
5. СТАТИСТИКА: проверено / нарушений / предупреждений / не проверено
6. ОБЩИЙ ВЫВОД

Кратко, деловым языком, формат официального документа."""

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

app = Flask(__name__)


# ============================================================
# MESSENGER BOT API
# ============================================================

def messenger_headers(content_type=None) -> dict:
    headers = {"Authorization": f"OAuth {MESSENGER_BOT_TOKEN}"}
    if content_type:
        headers["Content-Type"] = content_type
    return headers


def send_text(text: str, chat_id: str = None, login: str = None,
              reply_message_id: int = None):
    url = f"{MESSENGER_BASE_URL}/messages/sendText/"
    payload = {"text": text[:6000]}
    if chat_id:
        payload["chat_id"] = chat_id
    elif login:
        payload["login"] = login
    else:
        raise ValueError("Нужен chat_id или login")
    if reply_message_id:
        payload["reply_message_id"] = reply_message_id
    resp = requests.post(url, headers=messenger_headers("application/json"),
                         json=payload)
    resp.raise_for_status()
    data = resp.json()
    if not data.get("ok"):
        logger.error("sendText error: %s", data.get("description"))
    return data


def send_long_text(text: str, chat_id: str = None, login: str = None,
                   reply_message_id: int = None):
    chunks = [text[i:i + 6000] for i in range(0, len(text), 6000)]
    for i, chunk in enumerate(chunks):
        send_text(chunk, chat_id=chat_id, login=login,
                  reply_message_id=reply_message_id if i == 0 else None)


def send_file(file_path: str, chat_id: str = None, login: str = None):
    url = f"{MESSENGER_BASE_URL}/messages/sendFile/"
    data = {}
    if chat_id:
        data["chat_id"] = chat_id
    elif login:
        data["login"] = login
    else:
        raise ValueError("Нужен chat_id или login")
    file_name = os.path.basename(file_path)
    with open(file_path, "rb") as f:
        files = {"document": (file_name, f)}
        resp = requests.post(
            url,
            headers={"Authorization": f"OAuth {MESSENGER_BOT_TOKEN}"},
            data=data, files=files)
    resp.raise_for_status()
    return resp.json()


def get_file(file_id: str) -> bytes:
    url = f"{MESSENGER_BASE_URL}/messages/getFile/"
    resp = requests.get(url,
                        headers={"Authorization": f"OAuth {MESSENGER_BOT_TOKEN}"},
                        params={"file_id": file_id}, stream=True)
    resp.raise_for_status()
    ct = resp.headers.get("Content-Type", "")
    if "application/json" in ct:
        data = resp.json()
        if not data.get("ok"):
            raise RuntimeError(f"getFile error: {data.get('description')}")
    return resp.content


def get_updates(offset: int = 0, limit: int = 100) -> dict:
    url = f"{MESSENGER_BASE_URL}/messages/getUpdates/"
    resp = requests.get(url, headers=messenger_headers(),
                        params={"offset": offset, "limit": limit})
    resp.raise_for_status()
    return resp.json()


def set_webhook(webhook_url):
    url = f"{MESSENGER_BASE_URL}/self/update/"
    resp = requests.post(url, headers=messenger_headers("application/json"),
                         json={"webhook_url": webhook_url})
    resp.raise_for_status()
    return resp.json()


# ============================================================
# АУДИО
# ============================================================

def detect_audio_format(audio_bytes: bytes, file_name: str = "") -> str:
    header = audio_bytes[:16] if len(audio_bytes) >= 16 else audio_bytes
    if header[:4] == b'OggS':
        return "oggopus"
    if header[:3] == b'ID3' or (len(header) >= 2 and header[0] == 0xFF
                                 and (header[1] & 0xE0) == 0xE0):
        return "mp3"
    if header[:4] == b'RIFF':
        return "lpcm"
    if header[:4] in (b'fLaC', b'\x1a\x45\xdf\xa3'):
        return None
    if len(header) >= 8 and header[4:8] == b'ftyp':
        return None
    if header[:6] == b'#!AMR\n':
        return None
    ext = os.path.splitext(file_name)[1].lower() if file_name else ""
    return {".ogg": "oggopus", ".opus": "oggopus",
            ".mp3": "mp3", ".wav": "lpcm"}.get(ext)


def get_wav_sample_rate(audio_bytes: bytes) -> int:
    if len(audio_bytes) >= 28 and audio_bytes[:4] == b'RIFF':
        return struct.unpack_from('<I', audio_bytes, 24)[0]
    return 48000


def convert_to_ogg_opus(audio_bytes: bytes, file_name: str = "") -> bytes:
    ext = os.path.splitext(file_name)[1].lower() if file_name else ".bin"
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp_in:
        tmp_in.write(audio_bytes)
        tmp_in_path = tmp_in.name
    tmp_out_path = tmp_in_path + ".ogg"
    try:
        subprocess.run(
            [FFMPEG_PATH, "-y", "-i", tmp_in_path, "-acodec", "libopus",
             "-ac", "1", "-ar", "48000", "-b:a", "48k", tmp_out_path],
            capture_output=True, text=True, timeout=30, check=True)
        with open(tmp_out_path, "rb") as f:
            return f.read()
    finally:
        for p in (tmp_in_path, tmp_out_path):
            try: os.unlink(p)
            except OSError: pass


def split_audio_to_chunks(audio_bytes: bytes, file_name: str = "",
                          chunk_duration: int = 25) -> list:
    ext = os.path.splitext(file_name)[1].lower() if file_name else ".ogg"
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp_in:
        tmp_in.write(audio_bytes)
        tmp_in_path = tmp_in.name
    chunks = []
    try:
        probe = subprocess.run(
            [FFMPEG_PATH, "-i", tmp_in_path, "-f", "null", "-"],
            capture_output=True, text=True, timeout=30)
        duration_sec = 0
        for line in probe.stderr.split("\n"):
            if "Duration:" in line:
                parts = line.split("Duration:")[1].split(",")[0].strip()
                h, m, s = parts.split(":")
                duration_sec = int(h) * 3600 + int(m) * 60 + float(s)
                break
        if duration_sec <= chunk_duration:
            chunks.append(audio_bytes)
            return chunks
        start = 0
        idx = 0
        while start < duration_sec:
            tmp_out = tmp_in_path + f"_chunk{idx}.ogg"
            subprocess.run(
                [FFMPEG_PATH, "-y", "-i", tmp_in_path, "-ss", str(start),
                 "-t", str(chunk_duration), "-acodec", "libopus", "-ac", "1",
                 "-ar", "48000", "-b:a", "48k", tmp_out],
                capture_output=True, text=True, timeout=30)
            if os.path.exists(tmp_out):
                with open(tmp_out, "rb") as f:
                    data = f.read()
                if data:
                    chunks.append(data)
                try: os.unlink(tmp_out)
                except OSError: pass
            start += chunk_duration
            idx += 1
    finally:
        try: os.unlink(tmp_in_path)
        except OSError: pass
    return chunks


# ============================================================
# SPEECHKIT v1
# ============================================================

def recognize_speech_v1(audio_bytes: bytes, file_name: str = "") -> str:
    fmt = detect_audio_format(audio_bytes, file_name) or "oggopus"
    params = {"folderId": YC_FOLDER_ID, "lang": "ru-RU",
              "topic": "general", "format": fmt}
    if fmt == "lpcm":
        params["sampleRateHertz"] = str(get_wav_sample_rate(audio_bytes))
    resp = requests.post(SPEECHKIT_V1_URL,
                         headers={"Authorization": f"Api-Key {YC_API_KEY}"},
                         params=params, data=audio_bytes)
    if resp.status_code != 200:
        raise RuntimeError(f"SpeechKit error {resp.status_code}: {resp.text}")
    return resp.json().get("result", "")


def recognize_speech(audio_bytes: bytes, file_name: str = "") -> str:
    fmt = detect_audio_format(audio_bytes, file_name)
    if fmt is None:
        audio_bytes = convert_to_ogg_opus(audio_bytes, file_name)
        file_name = "converted.ogg"
    chunks = split_audio_to_chunks(audio_bytes, file_name, chunk_duration=25)
    texts = []
    for i, chunk in enumerate(chunks):
        try:
            t = recognize_speech_v1(chunk, "chunk.ogg")
            if t.strip():
                texts.append(t)
        except Exception as e:
            logger.warning("Chunk %d failed: %s", i + 1, e)
    return " ".join(texts)


# ============================================================
# YANDEX GPT
# ============================================================

def call_yandexgpt(system_prompt: str, user_message: str,
                   temperature: float = 0.3, max_tokens: int = 4000) -> str:
    headers = {"Content-Type": "application/json",
               "Authorization": f"Api-Key {YC_API_KEY}",
               "x-folder-id": YC_FOLDER_ID}
    payload = {
        "modelUri": YANDEXGPT_MODEL_URI,
        "completionOptions": {"stream": False, "temperature": temperature,
                              "maxTokens": max_tokens},
        "messages": [
            {"role": "system", "text": system_prompt},
            {"role": "user", "text": user_message},
        ],
    }
    resp = requests.post(YANDEXGPT_COMPLETION_URL, headers=headers, json=payload)
    if resp.status_code != 200:
        raise RuntimeError(f"YandexGPT error {resp.status_code}: {resp.text}")
    data = resp.json()
    try:
        return data["result"]["alternatives"][0]["message"]["text"]
    except (KeyError, IndexError):
        raise RuntimeError(f"Unexpected YandexGPT response: {data}")


def parse_json_from_llm(text: str) -> any:
    """Извлекает JSON из ответа LLM, даже если обёрнут в markdown."""
    # Убираем markdown-блоки
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    text = text.strip()
    return json.loads(text)


def process_with_llm(recognized_text: str) -> dict:
    results = {}

    # ШАГ 1: Извлечение
    logger.info("LLM Step 1: Extract")
    raw_extract = call_yandexgpt(
        SYSTEM_PROMPT_EXTRACT,
        f"Расшифровка:\n\n{recognized_text}",
        temperature=0.1, max_tokens=3000)
    try:
        extracted = parse_json_from_llm(raw_extract)
    except json.JSONDecodeError:
        logger.warning("Failed to parse extracted JSON, using raw text")
        extracted = {"notes": raw_extract, "findings": []}
    results["extracted"] = extracted

    # ШАГ 2: Заполнение чек-листа
    logger.info("LLM Step 2: Fill checklist")
    raw_checklist = call_yandexgpt(
        SYSTEM_PROMPT_FILL,
        f"Данные:\n\n{json.dumps(extracted, ensure_ascii=False)}",
        temperature=0.2, max_tokens=8000)
    try:
        checklist_items = parse_json_from_llm(raw_checklist)
    except json.JSONDecodeError:
        logger.warning("Failed to parse checklist JSON")
        checklist_items = []
    results["checklist_items"] = checklist_items

    # Строим словарь item_id → данные
    checklist_map = {}
    if isinstance(checklist_items, list):
        for item in checklist_items:
            iid = item.get("item_id", "")
            checklist_map[iid] = item
    results["checklist_map"] = checklist_map

    # ШАГ 3: Итоговый акт
    logger.info("LLM Step 3: Summary")
    summary = call_yandexgpt(
        SYSTEM_PROMPT_SUMMARY,
        f"Чек-лист:\n\n{raw_checklist}",
        temperature=0.2, max_tokens=3000)
    results["summary"] = summary

    return results


# ============================================================
# ГЕНЕРАЦИЯ DOCX
# ============================================================

# Цвета
COLOR_GREEN  = RGBColor(0x27, 0xAE, 0x60)
COLOR_RED    = RGBColor(0xE7, 0x4C, 0x3C)
COLOR_ORANGE = RGBColor(0xF3, 0x9C, 0x12)
COLOR_GRAY   = RGBColor(0x95, 0xA5, 0xA6)
COLOR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_DARK   = RGBColor(0x2C, 0x3E, 0x50)
COLOR_HEADER_BG = RGBColor(0x2C, 0x3E, 0x50)


def set_cell_shading(cell, color_hex: str):
    """Устанавливает фон ячейки таблицы."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.find(qn('w:shd'))
    if shading_elem is None:
        from lxml import etree
        shading_elem = etree.SubElement(shading, qn('w:shd'))
    shading_elem.set(qn('w:fill'), color_hex)
    shading_elem.set(qn('w:val'), 'clear')


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9,
                  color: RGBColor = None, alignment=None):
    """Заполняет ячейку текстом с форматированием."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def assessment_to_color(assessment: str) -> str:
    """Возвращает hex-цвет фона по оценке."""
    a = assessment.upper().strip()
    if a == "ДА":
        return "D5F5E3"   # зелёный
    elif a == "НЕТ":
        return "FADBD8"   # красный
    elif a == "Н/П":
        return "FCF3CF"   # жёлтый
    else:
        return "F2F3F4"   # серый — НЕ ПРОВЕРЕНО


def generate_checklist_docx(results: dict, recognized_text: str,
                            inspector_login: str = "") -> str:
    """Генерирует .docx файл с заполненным чек-листом."""

    doc = Document()

    # --- Стили ---
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)

    # --- Настройка страницы ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    extracted = results.get("extracted", {})
    checklist_map = results.get("checklist_map", {})
    summary = results.get("summary", "")
    now = datetime.now()

    # ============================================
    # ТИТУЛЬНАЯ СТРАНИЦА
    # ============================================

    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ЧЕК-ЛИСТ\nПРОМЫШЛЕННОЙ И ПРОИЗВОДСТВЕННОЙ\nБЕЗОПАСНОСТИ")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_DARK

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Сформирован автоматически на основе голосового отчёта")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Информация об объекте
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("Объект проверки", extracted.get("object_name", "—")),
        ("Адрес", extracted.get("address", "—")),
        ("Дата проверки", now.strftime("%d.%m.%Y")),
        ("Время проверки", now.strftime("%H:%M")),
        ("Проверяющий / Аудитор", extracted.get("inspector", inspector_login or "—")),
        ("Директор / Управляющий", extracted.get("director", "—")),
        ("Сотрудники на смене", extracted.get("employees_on_shift", "—")),
        ("Примечание", extracted.get("notes", "—")),
    ]

    for i, (label, value) in enumerate(info_data):
        set_cell_shading(info_table.cell(i, 0), "ECF0F1")
        set_cell_text(info_table.cell(i, 0), label, bold=True, size=10)
        set_cell_text(info_table.cell(i, 1), str(value) if value else "—", size=10)
        info_table.cell(i, 0).width = Cm(6)
        info_table.cell(i, 1).width = Cm(12)

    doc.add_page_break()

    # ============================================
    # БЛОКИ ЧЕК-ЛИСТА (21 блок)
    # ============================================

    total_checked = 0
    total_yes = 0
    total_no = 0
    total_na = 0
    total_not_checked = 0

    for block in CHECKLIST_BLOCKS:
        # Заголовок блока
        heading = doc.add_paragraph()
        run = heading.add_run(f"БЛОК {block['number']}. {block['title']}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_DARK

        # Таблица блока
        num_items = len(block["items"])
        table = doc.add_table(rows=num_items + 1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Ширины колонок
        for row in table.rows:
            row.cells[0].width = Cm(9)
            row.cells[1].width = Cm(3)
            row.cells[2].width = Cm(4)
            row.cells[3].width = Cm(2)

        # Заголовки
        headers = ["Критерий проверки", "Оценка", "Комментарий", "Срок устр."]
        for j, h in enumerate(headers):
            set_cell_shading(table.cell(0, j), "2C3E50")
            set_cell_text(table.cell(0, j), h, bold=True, size=9,
                         color=COLOR_WHITE,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Строки с критериями
        for i, item_text in enumerate(block["items"]):
            row_idx = i + 1
            # Извлекаем номер пункта
            item_id_match = re.match(r'^(\d+\.\d+)', item_text)
            item_id = item_id_match.group(1) if item_id_match else ""

            # Данные из LLM
            llm_data = checklist_map.get(item_id, {})
            assessment = llm_data.get("assessment", "НЕ ПРОВЕРЕНО")
            comment = llm_data.get("comment", "")
            deadline = llm_data.get("deadline", "")

            # Статистика
            a = assessment.upper().strip()
            if a == "ДА":
                total_yes += 1
                total_checked += 1
            elif a == "НЕТ":
                total_no += 1
                total_checked += 1
            elif a == "Н/П":
                total_na += 1
                total_checked += 1
            else:
                total_not_checked += 1

            # Критерий
            set_cell_text(table.cell(row_idx, 0), item_text, size=8)

            # Оценка с цветом
            bg = assessment_to_color(assessment)
            set_cell_shading(table.cell(row_idx, 1), bg)
            set_cell_text(table.cell(row_idx, 1), assessment, bold=True,
                         size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Комментарий
            set_cell_text(table.cell(row_idx, 2), comment or "", size=8)

            # Срок
            set_cell_text(table.cell(row_idx, 3), deadline or "", size=8,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph("")  # отступ между блоками

    doc.add_page_break()

    # ============================================
    # ИТОГОВАЯ СТАТИСТИКА
    # ============================================

    heading = doc.add_paragraph()
    run = heading.add_run("ИТОГИ ПРОВЕРКИ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_DARK

    total_items = total_yes + total_no + total_na + total_not_checked

    stats_table = doc.add_table(rows=6, cols=2)
    stats_table.style = 'Table Grid'

    stats = [
        ("Всего пунктов", str(total_items)),
        ("Соответствует (ДА)", str(total_yes)),
        ("Не соответствует (НЕТ)", str(total_no)),
        ("Неприменимо (Н/П)", str(total_na)),
        ("Не проверено", str(total_not_checked)),
        ("Процент выполнения",
         f"{(total_yes / total_checked * 100):.0f}%" if total_checked > 0 else "—"),
    ]

    colors_stats = ["ECF0F1", "D5F5E3", "FADBD8", "FCF3CF", "F2F3F4", "D6EAF8"]
    for i, ((label, value), bg) in enumerate(zip(stats, colors_stats)):
        set_cell_shading(stats_table.cell(i, 0), "ECF0F1")
        set_cell_text(stats_table.cell(i, 0), label, bold=True, size=10)
        set_cell_shading(stats_table.cell(i, 1), bg)
        set_cell_text(stats_table.cell(i, 1), value, bold=True, size=12,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")

    # ============================================
    # ИТОГОВЫЙ АКТ
    # ============================================

    heading = doc.add_paragraph()
    run = heading.add_run("ИТОГОВЫЙ АКТ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_DARK

    if summary:
        for line in summary.split("\n"):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ============================================
    # РАСШИФРОВКА АУДИО
    # ============================================

    heading = doc.add_paragraph()
    run = heading.add_run("ПРИЛОЖЕНИЕ: ИСХОДНАЯ РАСШИФРОВКА АУДИО")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_GRAY

    p = doc.add_paragraph(recognized_text)
    p.style.font.size = Pt(9)

    # ============================================
    # ПОДПИСИ
    # ============================================

    doc.add_paragraph("")
    doc.add_paragraph("")

    sign_table = doc.add_table(rows=2, cols=2)
    set_cell_text(sign_table.cell(0, 0), "Проверяющий:", bold=True, size=10)
    set_cell_text(sign_table.cell(0, 1),
                  f"{extracted.get('inspector', inspector_login)} _______________",
                  size=10)
    set_cell_text(sign_table.cell(1, 0), "Дата:", bold=True, size=10)
    set_cell_text(sign_table.cell(1, 1),
                  f"{now.strftime('%d.%m.%Y')} _______________", size=10)

    # ============================================
    # СОХРАНЕНИЕ
    # ============================================

    date_str = now.strftime("%Y%m%d_%H%M%S")
    obj_name = extracted.get("object_name", "object")
    # Очищаем имя от спецсимволов
    safe_name = re.sub(r'[^\w\s-]', '', obj_name)[:30].strip().replace(" ", "_")
    if not safe_name:
        safe_name = "checklist"

    file_path = os.path.join(
        tempfile.gettempdir(),
        f"checklist_{safe_name}_{date_str}.docx"
    )
    doc.save(file_path)
    logger.info("Generated DOCX: %s", file_path)
    return file_path


# ============================================================
# ОБРАБОТКА ОБНОВЛЕНИЙ
# ============================================================

def is_audio_file(file_name: str) -> bool:
    if not file_name:
        return False
    return os.path.splitext(file_name)[1].lower() in AUDIO_EXTENSIONS


def resolve_reply_target(update: dict) -> dict:
    chat = update.get("chat", {})
    chat_type = chat.get("type")
    from_user = update.get("from", {})
    if chat_type in ("group", "channel"):
        return {"chat_id": chat["id"]}
    login = from_user.get("login")
    if login:
        return {"login": login}
    return {}


def process_update(update: dict):
    logger.info("Processing update_id=%s from=%s",
                update.get("update_id"),
                update.get("from", {}).get("login", "?"))

    reply_to = resolve_reply_target(update)
    if not reply_to:
        return

    message_id = update.get("message_id")
    from_login = update.get("from", {}).get("login", "unknown")

    # ─── ФАЙЛ ───────────────────────────────────────────────
    file_info = update.get("file")
    if file_info:
        file_id = file_info["id"]
        file_name = file_info.get("name", "")

        if not is_audio_file(file_name):
            send_text(f"📎 «{file_name}» — не аудио. Отправьте голосовое.",
                      **reply_to)
            return

        send_text("⏳ Принял аудио. Начинаю обработку…", **reply_to)

        try:
            audio_bytes = get_file(file_id)

            send_text("🎤 Распознаю речь…", **reply_to)
            recognized_text = recognize_speech(audio_bytes, file_name)

            if not recognized_text.strip():
                send_text("🤷 Не удалось распознать речь.",
                          reply_message_id=message_id, **reply_to)
                return

            send_long_text(f"📝 Расшифровка:\n\n{recognized_text}", **reply_to)

            send_text("🤖 Анализирую и заполняю чек-лист (3 этапа)…", **reply_to)
            results = process_with_llm(recognized_text)

            summary = results.get("summary", "")
            if summary:
                send_long_text(f"📋 ИТОГОВЫЙ АКТ\n\n{summary}",
                              reply_message_id=message_id, **reply_to)

            send_text("📄 Формирую документ DOCX…", **reply_to)
            docx_path = generate_checklist_docx(
                results, recognized_text, from_login)
            send_file(docx_path, **reply_to)
            send_text("✅ Готово! Чек-лист отправлен файлом.", **reply_to)

            try: os.unlink(docx_path)
            except OSError: pass

        except Exception as e:
            logger.exception("Error: %s", e)
            send_text(f"❌ Ошибка: {e}", **reply_to)
        return

    # ─── ТЕКСТ ──────────────────────────────────────────────
    text = update.get("text", "")
    if text:
        cmd = text.strip().lower()
        if cmd in ("/start", "/help", "помощь"):
            send_text(
                "🎙 Бот проверки промышленной безопасности\n\n"
                "Отправьте голосовое сообщение с описанием обхода объекта.\n\n"
                "Бот:\n"
                "1️⃣ Распознает речь (SpeechKit)\n"
                "2️⃣ Извлечёт данные (YandexGPT)\n"
                "3️⃣ Заполнит чек-лист (21 блок, 100+ пунктов)\n"
                "4️⃣ Сформирует итоговый акт\n"
                "5️⃣ Отправит DOCX-документ с таблицами\n\n"
                "/help — справка\n"
                "/blocks — список блоков",
                **reply_to)
        elif cmd == "/blocks":
            blocks = "\n".join(
                f"{b['number']}. {b['title']}" for b in CHECKLIST_BLOCKS)
            send_text(f"📋 Блоки чек-листа (21):\n\n{blocks}", **reply_to)
        else:
            send_text("🎤 Отправьте голосовое сообщение.\n/help — справка",
                      **reply_to)
        return

    if update.get("sticker") or update.get("images"):
        send_text("🎤 Я принимаю только аудио.", **reply_to)


# ============================================================
# WEBHOOK / POLLING / MAIN
# ============================================================

@app.route("/webhook", methods=["POST"])
def webhook_handler():
    data = request.get_json(force=True)
    if not data.get("ok"):
        return jsonify({"status": "error"}), 200
    for upd in data.get("updates", []):
        try: process_update(upd)
        except Exception as e: logger.exception("Error: %s", e)
    return jsonify({"status": "ok"}), 200


@app.route("/health", methods=["GET"])
def health():
    return "OK", 200


def run_polling():
    logger.info("Starting polling…")
    offset = 0
    while True:
        try:
            data = get_updates(offset=offset, limit=100)
            if not data.get("ok"):
                time.sleep(5)
                continue
            for upd in data.get("updates", []):
                try: process_update(upd)
                except Exception as e: logger.exception("Error: %s", e)
                uid = upd.get("update_id", 0)
                if uid >= offset:
                    offset = uid + 1
            if not data.get("updates"):
                time.sleep(1)
        except requests.exceptions.RequestException as e:
            logger.error("Network: %s", e)
            time.sleep(5)
        except Exception as e:
            logger.exception("Unexpected: %s", e)
            time.sleep(5)


if __name__ == "__main__":
    import sys
    mode = sys.argv[1] if len(sys.argv) > 1 else "polling"
    if mode == "webhook":
        webhook_url = os.environ.get("WEBHOOK_URL", "https://example.com/webhook")
        set_webhook(webhook_url)
        app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
    elif mode == "polling":
        try: set_webhook(None)
        except: pass
        run_polling()
